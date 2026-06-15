# -*- coding: utf-8 -*-
import docx, os, sys
from docx.document import Document as _Doc
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

base = r'C:\Users\ADELEKEOLORUNISOLAO\Desktop\extracted-vocabs - Copy'
out = open(os.path.join('scripts', 'docx_dump.txt'), 'w', encoding='utf-8')

def iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

for f in ['Goethe_Vocabulary_A1.docx','Goethe_Vocabulary_A2.docx','Goethe_Vocabulary_B1.docx','Goethe_Vocabulary_B2.docx']:
    d = docx.Document(os.path.join(base, f))
    out.write('\n' + '='*70 + '\n' + f + '\n' + '='*70 + '\n')
    out.write('paragraphs=%d tables=%d\n' % (len(d.paragraphs), len(d.tables)))
    # Document order: show sequence of headings + tables
    seq = list(iter_block_items(d))
    out.write('block items=%d\n' % len(seq))
    tbl_count = 0
    for i, b in enumerate(seq):
        if isinstance(b, Paragraph):
            t = b.text.strip()
            if t:
                out.write('P[%s]: %s\n' % (b.style.name, t[:140]))
        else:
            tbl_count += 1
            rows = b.rows
            ncols = len(b.columns)
            out.write('  TABLE #%d rows=%d cols=%d\n' % (tbl_count, len(rows), ncols))
            # header + up to 4 sample rows
            for r in rows[:5]:
                cells = [c.text.strip().replace('\n',' / ') for c in r.cells]
                out.write('     | ' + ' || '.join(cells) + '\n')
            if f != 'Goethe_Vocabulary_A1.docx':
                # for bigger files only show first 6 tables to keep dump small
                if tbl_count >= 6:
                    out.write('  ... (remaining tables omitted in dump)\n')
                    break
    out.flush()
out.close()
print('done')
