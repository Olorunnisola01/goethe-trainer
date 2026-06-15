#!/usr/bin/env python3
"""
Extract vocabulary lists from the Goethe Trainer website data (public/data/vocab.json)
into separate professional .docx files for A1, A2, B1, B2 levels.

Only uses the vocab data served by the website itself.
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, fill_color: str):
    """Set background color for a table cell (e.g. '4472C4' for blue header)."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=60, bottom=40, left=60, right=60):
    """Set cell padding (twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_vocab_document(level: str, categories: list, output_path: Path):
    """Create a nicely formatted DOCX for one CEFR level."""
    doc = Document()

    # A4 page (common for European/German language materials), narrow margins for lists
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title = doc.add_heading(f"Goethe Trainer — {level} Vocabulary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)  # deep blue

    # Subtitle / source note
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Extracted exclusively from the Goethe Trainer website vocabulary data")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source.add_run("Source: public/data/vocab.json (as served by the live web application)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Summary stats
    total_entries = sum(len(cat.get("entries", [])) for cat in categories)
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats.add_run(f"📚 {len(categories)} categories  •  {total_entries} vocabulary items / phrases")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()  # visual spacer

    # Color per level (subtle header accent)
    level_colors = {
        "A1": "15803d",  # green
        "A2": "1d4ed8",  # blue
        "B1": "7c3aed",  # violet
        "B2": "b45309",  # amber
    }
    header_fill = level_colors.get(level, "2E75B6")

    # Per-category sections
    for idx, cat in enumerate(categories, 1):
        emoji = cat.get("emoji", "📖")
        title_text = cat.get("title", "UNTITLED").upper()

        # Category heading (H2)
        h = doc.add_heading(f"{emoji}  {title_text}", level=2)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)

        entries = cat.get("entries", [])
        if not entries:
            continue

        # Table: German | English | Example DE | Example EN
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # type: ignore[attr-defined]

        # Header row
        header_cells = table.rows[0].cells
        headers = ["German", "English", "Example (German)", "Example (English)"]
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            set_cell_shading(cell, header_fill)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_margins(cell, top=40, bottom=30, left=50, right=50)

        # Data rows
        for entry in entries:
            row = table.add_row()
            cells = row.cells

            w = entry.get("w", "").strip()
            t = entry.get("t", "").strip()
            de = entry.get("de", "").strip()
            en = entry.get("en", "").strip()

            cells[0].text = w
            cells[1].text = t
            cells[2].text = de
            cells[3].text = en

            # Style data cells
            for i, cell in enumerate(cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Calibri"
                        if i == 0:  # German word column - make prominent
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)
                set_cell_margins(cell, top=30, bottom=25, left=50, right=50)

        # Reasonable column widths for A4 with 1.2cm margins (~17.5cm usable width)
        # German | English | DE ex | EN ex
        col_widths = [Cm(3.8), Cm(4.2), Cm(5.0), Cm(5.0)]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

        # Small spacer after table (except last)
        if idx < len(categories):
            doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—  End of vocabulary list for " + level + "  —")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This extraction contains only the vocabulary items displayed in the Goethe Trainer web application. "
        "Examples and translations are as provided in the site data."
    )
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Document core properties
    core_props = doc.core_properties
    core_props.author = "Goethe Trainer Vocabulary Extractor"
    core_props.title = f"Goethe Trainer {level} Vocabulary"
    core_props.subject = f"Complete A1–B2 level German vocabulary for {level} (website data only)"

    doc.save(str(output_path))
    print(f"[OK] Created {output_path.name}  ({len(categories)} cats, {total_entries} entries)")


def main():
    src = Path("public/data/vocab.json")
    if not src.exists():
        raise FileNotFoundError(f"Source not found: {src}")

    print(f"Loading vocabulary data from {src} ...")
    with open(src, "r", encoding="utf-8") as f:
        raw_data = json.load(f)

    by_level: dict[str, list] = {"A1": [], "A2": [], "B1": [], "B2": []}
    for cat in raw_data:
        lvl = cat.get("level")
        if lvl in by_level:
            by_level[lvl].append(cat)

    # Summary (only categories that contain actual vocabulary entries)
    print("\nVocabulary summary from website data (public/data/vocab.json):")
    for lvl in ["A1", "A2", "B1", "B2"]:
        cats = [c for c in by_level[lvl] if (c.get("entries") or [])]
        n_cats = len(cats)
        n_entries = sum(len(c.get("entries", [])) for c in cats)
        print(f"  {lvl}: {n_cats:3d} categories, {n_entries:4d} entries")

    out_dir = Path("extracted-vocabs")
    out_dir.mkdir(exist_ok=True)
    print(f"\nOutput directory: {out_dir.resolve()}")

    for level in ["A1", "A2", "B1", "B2"]:
        out_file = out_dir / f"Goethe_Vocabulary_{level}.docx"
        create_vocab_document(level, by_level[level], out_file)

    print("\n[OK] All vocabulary DOCX files generated successfully.")
    print("   Only data from the website (public/data/vocab.json) was used.")


if __name__ == "__main__":
    main()
